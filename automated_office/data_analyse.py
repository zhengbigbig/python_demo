import pandas as pd
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches  # 定义插入图片的英寸

os.chdir('/Users/zhengzhiheng/PycharmProjects/untitled3/myproject/automated_office')

imgname = 'test.jpg'
students = pd.read_excel('people.xlsx')

students.sort_values(by='Score', inplace=True, ascending=False)
plt.bar(students.Name, students.Score, color='orange')
plt.title('Student Score', fontsize=16)
plt.xlabel('Name')
plt.ylabel('Score')
plt.xticks(students.Name, rotation='45')
plt.tight_layout()
plt.savefig(imgname)

document = Document()
document.add_heading('数据分析报告', level=0)
first_student = students.iloc[0, :]['Name']
first_score = students.iloc[0, :]['Score']

p = document.add_paragraph('分数挂的第一的学生是：')
p.add_run(str(first_student)).bold = True
p.add_run(',分数为')
p.add_run(str(first_score)).bold = True

p1 = document.add_paragraph(f'总共有{len(students.Name)}名学生参加了考试，学生考试的总体情况：')

table = document.add_table(rows=len(students.Name) + 1, cols=2)
table.style = 'LightShading-Accent1'
table.cell(0, 0).text = '学生姓名'
table.cell(0, 1).text = '学生分数'
for i, (index, row) in enumerate(students.iterrows()):
	table.cell(i + 1, 0).text = str(row['Name'])
	table.cell(i + 1, 1).text = str(row['Score'])

document.add_picture(imgname)
document.save('Students.docx')
print('Done')
