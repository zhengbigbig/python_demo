from docx import Document
from docx.shared import Inches # 定义插入图片的英寸
import os

os.chdir('/Users/zhengzhiheng/PycharmProjects/untitled3/myproject/automated_office')

document = Document()
# 新增段落
p2 = document.add_paragraph('这是一个段落')
# 在p2段落插入一个段落
p1 = p2.insert_paragraph_before('第一个段落')
# 新增标题,level 1-9
# level=0，新增的标签会有下划线样式
document.add_heading('这是标题', level=1)
document.add_page_break()  # 分页符
# 新增表格，因为前面有分页符，所有这里的内容在新的一页
table = document.add_table(rows=6, cols=6)
# 第一行第二列
cell = table.cell(0, 2)
cell.text = '第一行第二列'
row = table.rows[1]
row.cells[0].text = '第二行第一列'
row.cells[1].text = '第二行第二列'
# 添加图片
document.add_picture('header.jpg', width=Inches(1.25))
document.save('new.docx')
